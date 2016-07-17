'''
Create Agenda
generate agenda from Asana tasks
and place them in Alpha Alpha Gdrive
folder to share with the chapter.

author(s):
	Jorge Rojas
'''

'''Utility Imports'''
import sys
import time
import json, base64
from datetime import datetime
from utilities import curl, info
from optparse import OptionParser

'''gDrive Imports'''
from docx import Document
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

today = lambda: datetime.fromtimestamp( time.time() )
#Change '***************:' auth token from Asana once the members are registered
auth_token = base64.encodestring( '3IOlvKIK.qCLaoBd9o9vQzENWfspMQUl:' ).replace( '\n', '' )

heading = 'La Unidad Latina, Lambda Upsilon Lambda Fraternity, Inc.'
createDate = 'Alpha Alpha Chapter Agenda {0.year}-{0.month}-{0.day}'.format( today() )
attendance = '{0} - {1}'	#1 - Name; 2 - Status
topic_section = '{0} - mediator: {1}'	#discussion topic heading: 0 - topic number; 1 - topic name;

def getUser( taskStructDict ):
	user_id = taskStructDict[ 'assignee' ][ 'id' ]
	user_name = taskStructDict[ 'assignee' ][ 'name' ]

	user_url = 'https://app.asana.com/api/1.0/users/{0}?opt_pretty'.format(user_id)
	user = curl( user_url, authToken = auth_token )
	j_user = json.loads( user )

	return j_user[ 'data' ][ 'name' ], j_user[ 'data' ][ 'email' ]

def postLinkToSlack( gFileLink, chan = 'debug' ):
	slack_webhook_url = 'https://hooks.slack.com/services/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'
	slack_msg = 'Hey <!everyone>, <{0}|click this link> to see this week\'s agenda!'.format(gFileLink)
	slack_payload = '{"text": "%s","channel": "#%s","username": "ReminderBot","icon_emoji": ":mega:"}'%(slack_msg, chan)
	# print slack_payload
	slack_res = curl( slack_webhook_url, slack_payload )
	print "Agenda's link was succesfully posted on Slack."

def uploadToDrive( drive, fileName, gFolderID, slackChan):
	'''Google Drive Authentication'''
	if drive == None:
		gauth = GoogleAuth()
		gauth.LocalWebserverAuth()
		drive = GoogleDrive(gauth)	#gdrive service instance

	#fileName[8:] removes agendas/ from filename
	#agenda/ used for local storage
	try:
		drive_agenda = drive.CreateFile({'title': fileName[8:], 'parents': [{"id": gFolderID}]})
		drive_agenda.SetContentFile( fileName )
		drive_agenda.Upload()
		print 'Agenda was successfully saved to Google Drive!'
		# else:
		# 	drive_agenda = drive.CreateFile({'title': fileName[8:], 'parents': [{"id": parent_folder_id}]})
	except:
		print 'Agenda was successfully created, but failed to upload to the drive!'

	g_alternate_link = ''
	files = drive.ListFile({'q': "'{0}' in parents and trashed=false".format(gFolderID)}).GetList()
	
	for cur_file in files:
		if cur_file['title'] == fileName[8:]:
			g_alternate_link = cur_file['alternateLink']

	postLinkToSlack( g_alternate_link, slackChan )

def writeAgenda( authToken, projectUrl, document, fileName ):
	'''Capture all tasks (discussion topics)'''
	all_tasks = curl( projectUrl, authToken = authToken )
	j_all_tasks = json.loads( all_tasks )

	topic_cnt = 1
	isAttendance = False

	for cur_task in j_all_tasks[ 'data' ]:
		
		cur_task_name = cur_task[ 'name' ]
		cur_task_id = cur_task[ 'id' ]

		if cur_task_name.endswith( ':' ):
			if cur_task_name.find( 'Attendance') >= 0:
				isAttendance = True
			else:
				isAttendance = False
			document.add_heading( cur_task_name[:-1], level = 1 )
			topic_cnt = 1
		else:
			task_url = 'https://app.asana.com/api/1.0/tasks/{0}'.format(cur_task_id)
			story_url = 'https://app.asana.com/api/1.0/tasks/{0}/stories'.format(cur_task_id)

			'''Task Struct'''
			cur_task_struct = curl( task_url, authToken = authToken )
			j_cur_task_struct = json.loads( cur_task_struct )
			j_cur_task = j_cur_task_struct[ 'data' ]

			'''Stories Struct'''
			cur_task_stories = curl( story_url, authToken = authToken )
			j_cur_task_stories = json.loads( cur_task_stories )
			stories = j_cur_task_stories[ 'data' ]	#list of dictionaries 

			if not j_cur_task[ 'completed' ] and isAttendance:
				document.add_paragraph( attendance.format( j_cur_task[ 'assignee' ][ 'name' ], cur_task_name ), style = 'List Bullet')
			elif not j_cur_task[ 'completed' ] and not isAttendance:
				if j_cur_task['assignee'] != None and cur_task_name != None:
					document.add_heading( topic_section.format( cur_task_name.encode('utf-8'), j_cur_task[ 'assignee' ][ 'name' ].encode('utf-8') ).decode('utf-8'), level = 2)
				elif j_cur_task['assignee'] == None and cur_task_name != None:
					document.add_heading( topic_section.format( cur_task_name.encode('utf-8'), 'EBOARD' ).decode('utf-8'), level = 2)

				for cur_story in stories:
		
					tab_cnt = 0
					comments = cur_story[ 'text' ].encode('utf-8').split('\n')
					for cur_comment in comments:
						for cur_char in cur_comment:
							if cur_char != '_':
								break
							else:
								tab_cnt += 1

						comment = cur_comment[tab_cnt:].decode('utf-8')
						# print comment
						tab_cnt += 1
						if tab_cnt <= 1 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet')
						elif tab_cnt <= 3 and tab_cnt > 1 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet {0}'.format(tab_cnt) )
						elif tab_cnt > 3 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet 3')
						tab_cnt = 0

			topic_cnt += 1

	document.save( fileName )

def createAgenda(projKey, slackChan = None, drive = None):
	#Needs updating if path changes
	# gdrive_path = '00 Internal/03 Agendas & Minutes/2015 - 2016'.split('/')	
	info('function: createAgenda')
	projects = {
			'debug': '54999242167362', 
			'agenda': '47058473474991'
	}
	gFolders = {
			'debug': '0B2_x5hkUMMyZLW82bFo5dzhqdjQ', 
			'agenda':'0B2_x5hkUMMyZfl9mOHcwVEJqQnk5cjdwNnFNTTJSRUg1N3F0X1FyM0dGVVlBVXdVd25PMWc' 
	}
	project_url = 'https://app.asana.com/api/1.0/projects/{0}/tasks'.format(projects[projKey])

	'''Document Formatting Variables'''
	fi_name = 'agendas/lulaa_weekly_agenda_{0.year}-{0.month}-{0.day}.docx'.format( today() )
	# if proj_key == 'debug':
	# 	fi_name = 'lulaa_agenda_debug.docx'

	agenda_doc = Document()
	agenda_doc.add_heading( heading, 0 )
	agenda_doc.add_paragraph( createDate ).bold = True
	writeAgenda( auth_token, project_url, agenda_doc, fi_name )
	uploadToDrive( drive, fi_name, gFolders[projKey], slackChan )

if __name__ == '__main__':
	parser = OptionParser()

	parser.add_option( "--projKey",
					dest="projKey",
					default = "debug",
					help="Asana Project Key: debug or agenda. Default value is debug")

	parser.add_option( "--chan",
				dest="chan",
				default = "general",
				help="Slack channel name. Do not include the pound (#) sign. Default value is debug")

	(options, args) = parser.parse_args()
	proj_key = options.projKey	#Proj Key from terminal
	chan = options.chan
	# print proj_key
	createAgenda(proj_key, slackChan = chan, drive = None)

'''
TODO:
- Change Folder to save agendas from UI with shareable link
'''