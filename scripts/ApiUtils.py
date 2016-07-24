'''
API Utility Functions

Author(s):
	Jorge Rojas

Date:
	7/23/2016
'''

'''Utility Imports'''
import sys
import json
import urllib
import Utils as ut
from datetime import datetime
from optparse import OptionParser

''''''''''''''''''''''''
'''UPDATE ASANA TASKS'''
''''''''''''''''''''''''
slack_members = {}
receivers = []
receiver_names = []

def getSlackUsers( chapter ):
	user_list_url = 'https://slack.com/api/users.list?token={0}'.format(chapter['integrations']['slack']['auth_token'])
	slack_users = ut.curl( user_list_url )
	j_slack_users = json.loads( slack_users )
	# print j_slack_users
	for member in j_slack_users['members']:
		slack_members[ member['profile']['email'] ] = [ member['id'], member['name'] ]

def getAsanaUser( taskStructDict, chapter ):

	asana_auth_token = ut.encodeAuthToken(chapter['integrations']['asana']['auth_token'])
	user_id = taskStructDict[ 'assignee' ][ 'id' ]
	user_url = 'https://app.asana.com/api/1.0/users/{0}'.format(user_id)
	user = ut.curl( user_url, authToken = asana_auth_token )
	j_user = json.loads( user )

	return j_user[ 'data' ][ 'name' ], j_user[ 'data' ][ 'email' ]

def sendSlackReminder( userName, userEmail, taskName, channel, chapter):
	'''
	Remind the user via Slack mention
	'''

	member_info = slack_members[userEmail]
	slack_webhook = chapter['integrations']['slack']['webhook_url']
	slack_reminder_msg = '"Hey <@%s|%s>, this Asana task is pending: %s"'%( member_info[0], member_info[1], taskName )
	
	try:
		slack_res = ut.curl( slack_webhook, '{"channel": "#' + channel + '", "username": "ReminderBot", "text":' + slack_reminder_msg + ', "icon_emoji": ":mega:"}')
		print 'Successfully sent reminder to {0}.'.format(userName)
	except:
		print 'Failed to send reminder to {0}.'.format(userName)

def sendAsanaReminder( taskId, chapter):
	asana_auth_token = ut.encodeAuthToken(chapter['integrations']['asana']['auth_token'])
	story_url = 'https://app.asana.com/api/1.0/tasks/%s/stories'%taskId
	data = urllib.urlencode( { 'text': 'Please provide an update regarding this task and update the due date accordingly!' } )
	
	try:
		task_story = ut.curl( story_url, data, asana_auth_token )
		j_task_story = json.loads( task_story )[ 'data' ]
		print 'POST to Asana Successful.'
	except:
		print 'POST to Asana Failed.'

def remindUser( channel, taskStructDict, chapter ):

	global receivers, receiver_names

	task_id = taskStructDict[ 'id' ]
	task_name = taskStructDict[ 'name' ]
	user_name, user_email = getAsanaUser( taskStructDict, chapter )

	try:
		reminder_pref = chapter['members'][user_name]['reminder_pref'].split(',')	#assuming user name is same as name in mongoDB
	except KeyError:
		reminder_pref = {}
		print "{0} is not in our database. Please update your member list.".format(user_name)
	# receivers.append( str( user_email ) )
	# receiver_names.append( str( user_name ) )
	# print 'Name: %s >> Email: %s'%( user_name, user_email )	#DEBUG

	if 'asana' in reminder_pref:
		sendAsanaReminder(task_id, chapter)
	if 'slack' in reminder_pref:
		sendSlackReminder( user_name, user_email, task_name, channel, chapter )
	if 'twilio' in reminder_pref:
		sendSMSReminder(task_name, chapter['members'][user_name], chapter['integrations']['twilio'])
	if 'email' in reminder_pref:
		#Not yet implemented
		pass

def getAsanaTasks( channel, projectUrl, chapter ):
	asana_auth_token = ut.encodeAuthToken(chapter['integrations']['asana']['auth_token'])
	all_tasks = ut.curl( projectUrl, authToken = asana_auth_token )
	j_all_tasks = json.loads( all_tasks )
	
	for cur_task in j_all_tasks[ 'data' ]:
		'''
		Current Task >> dict
		keys:
			task id, name, due_on, assignee, completed, etc...
		'''
		#Filter out labels or sections
		cur_task_due = cur_task['due_on']
		cur_task_assignee = cur_task['assignee']
		cur_task_completed = cur_task['completed']

		if not cur_task_completed and cur_task_assignee is not None and cur_task_due is not None:
			#Converting delta time to milliseconds from current time
			cur_task_delta = ut.getDeltaTimeMilli(cur_task_due)
			# print cur_task_delta #DEBUG

			if ut.taskPending( cur_task_delta ):
				remindUser( channel, cur_task, chapter )

def updateAsana(projKey, slackChan, chapter):
	ut.info('function: updateAsana')

	projects = {'debug': '54999242167362', 'deliverables': '24426061282606'}
	project_url = 'https://app.asana.com/api/1.0/projects/%s/tasks?opt_fields=due_on,assignee,name,completed'%projects[ projKey ]
	
	getSlackUsers( chapter )
	getAsanaTasks( slackChan, project_url, chapter )

''''''''''''''''''''''''
'''CREATE AGENDA ASANA'''
''''''''''''''''''''''''
heading = 'La Unidad Latina, Lambda Upsilon Lambda Fraternity, Inc.'
createDate = 'Alpha Alpha Chapter Agenda {0.year}-{0.month}-{0.day}'.format( ut.timeNow() )
attendance = '{0} - {1}'	#1 - Name; 2 - Status
topic_section = '{0} - mediator: {1}'	#discussion topic heading: 0 - topic number; 1 - topic name;

def shareAgenda( gFileLink, chan = 'debug', slackWebhook = None ):
	slack_msg = 'Hey <!everyone>, <{0}|click this link> to see this week\'s agenda!'.format(gFileLink)
	slack_payload = '{"text": "%s","channel": "#%s","username": "ReminderBot","icon_emoji": ":mega:"}'%(slack_msg, chan)
	# print slack_payload
	slack_res = ut.curl( slackWebhook, slack_payload )
	print "Agenda's link was succesfully posted on Slack."

def uploadToDrive( drive, fileName, slackChan, chapter):
	'''Google Drive Authentication'''
	
	#Extract folder ID from link
	g_folder_id = chapter['gfolder_link'].split('=')[1]
	slack_webhook = chapter['integrations']['slack']['webhook_url']
	#fileName[8:] removes agendas/ from filename
	#agenda/ used for local storage
	try:
		drive_agenda = drive.CreateFile({'title': fileName[8:], 'parents': [{"id": g_folder_id}]})
		drive_agenda.SetContentFile( fileName )
		drive_agenda.Upload()
		print 'Agenda was successfully saved to Google Drive!'
	except:
		print 'Agenda was successfully created, but failed to upload to the drive!'

	g_alternate_link = ''
	files = drive.ListFile({'q': "'{0}' in parents and trashed=false".format(g_folder_id)}).GetList()
	
	for cur_file in files:
		if cur_file['title'] == fileName[8:]:
			g_alternate_link = cur_file['alternateLink']

	shareAgenda( g_alternate_link, slackChan, slack_webhook )

def writeAgenda( authToken, projectUrl, document, fileName ):
	'''Capture all tasks (discussion topics)'''
	auth_token = ut.encodeAuthToken(authToken)
	all_tasks = ut.curl( projectUrl, authToken = auth_token )
	j_all_tasks = json.loads( all_tasks )

	topic_cnt = 1
	isAttendance = False

	for cur_task in j_all_tasks[ 'data' ]:
		
		cur_task_name = cur_task[ 'name' ]
		cur_task_id = cur_task[ 'id' ]

		if cur_task_name.endswith( ':' ):
			if cur_task_name.find( 'Attendance') >= 0:
				isAttendance = True
			else:
				isAttendance = False
			document.add_heading( cur_task_name[:-1], level = 1 )
			topic_cnt = 1
		else:
			task_url = 'https://app.asana.com/api/1.0/tasks/{0}'.format(cur_task_id)
			story_url = 'https://app.asana.com/api/1.0/tasks/{0}/stories'.format(cur_task_id)

			'''Task Struct'''
			cur_task_struct = ut.curl( task_url, authToken = auth_token )
			j_cur_task_struct = json.loads( cur_task_struct )
			j_cur_task = j_cur_task_struct[ 'data' ]

			'''Stories Struct'''
			cur_task_stories = ut.curl( story_url, authToken = auth_token )
			j_cur_task_stories = json.loads( cur_task_stories )
			stories = j_cur_task_stories[ 'data' ]	#list of dictionaries 

			if not j_cur_task[ 'completed' ] and isAttendance:
				document.add_paragraph( attendance.format( j_cur_task[ 'assignee' ][ 'name' ], cur_task_name ), style = 'List Bullet')
			elif not j_cur_task[ 'completed' ] and not isAttendance:
				if j_cur_task['assignee'] != None and cur_task_name != None:
					document.add_heading( topic_section.format( cur_task_name.encode('utf-8'), j_cur_task[ 'assignee' ][ 'name' ].encode('utf-8') ).decode('utf-8'), level = 2)
				elif j_cur_task['assignee'] == None and cur_task_name != None:
					document.add_heading( topic_section.format( cur_task_name.encode('utf-8'), 'EBOARD' ).decode('utf-8'), level = 2)

				for cur_story in stories:
		
					tab_cnt = 0
					comments = cur_story[ 'text' ].encode('utf-8').split('\n')
					for cur_comment in comments:
						for cur_char in cur_comment:
							if cur_char != '_':
								break
							else:
								tab_cnt += 1

						comment = cur_comment[tab_cnt:].decode('utf-8')
						# print comment
						tab_cnt += 1
						if tab_cnt <= 1 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet')
						elif tab_cnt <= 3 and tab_cnt > 1 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet {0}'.format(tab_cnt) )
						elif tab_cnt > 3 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet 3')
						tab_cnt = 0

			topic_cnt += 1

	document.save( fileName )

def authDrive(gAuthCode = None):
	from pydrive.auth import GoogleAuth
	from pydrive.drive import GoogleDrive

	gauth = GoogleAuth()
	if gAuthCode == None:
		gauth.LocalWebserverAuth()
		drive = GoogleDrive(gauth)	#gdrive service instance
	else:
		gauth.Auth(gAuthCode)
		drive = GoogleDrive(gauth)

	return drive

def createAgenda(projKey, slackChan = None, gAuthCode = None, chapter = None):
	#Needs updating if path changes
	from docx import Document
	

	ut.info('function: createAgenda')
	projects = {
			'debug': '54999242167362', 
			'agenda': '47058473474991'
	}
	project_url = 'https://app.asana.com/api/1.0/projects/{0}/tasks'.format(projects[projKey])

	drive = authDrive(gAuthCode)
	fi_name = 'agendas/lulaa_weekly_agenda_{0.year}-{0.month}-{0.day}.docx'.format( ut.timeNow() )
	asana_auth_token = chapter['integrations']['asana']['auth_token']

	agenda_doc = Document()
	agenda_doc.add_heading( heading, 0 )
	agenda_doc.add_paragraph( createDate ).bold = True
	writeAgenda( asana_auth_token, project_url, agenda_doc, fi_name )
	uploadToDrive( drive, fi_name, slackChan, chapter )

''''''''''''''''''''''''
'''TWILIO   REMINDERS'''
''''''''''''''''''''''''
def sendSMSReminder(taskName, member, twilioTokens):
	from twilio.rest import TwilioRestClient

	if twilioTokens['accnt_sid'] is None or twilioTokens['auth_token'] is None:
		print 'Twilio tokens missing.'

	try:
		twilio_client = TwilioRestClient(twilioTokens['accnt_sid'], twilioTokens['auth_token'])
	except:
		print 'Twilio tokens are incorrect. Please check within '

	try:
		twilio_client.messages.create(
	        to = member['phone_num'],
	        from_ = twilioTokens['twilio_num'],
	        body = 'A friendly reminder that you have a certain task due: {0}.\n\nPlease provide an update to your VP when you get the chance'.format(taskName),
	    )
	except:
		print 'Message failed to send. Please check member and chapter phone numbers.'

if __name__ == '__main__':
	parser = OptionParser()

	parser.add_option( "--projKey",
					dest="projKey",
					default = "debug",
					help="Asana Project Key: debug or deliverables. Default value is debug")

	parser.add_option( "--chan",
					dest="chan",
					default = "debug",
					help="Slack channel name. Do not include the pound (#) sign. Default value is debug")

	(options, args) = parser.parse_args()
	projKey = options.projKey
	chan = options.chan
	updateAsana(projKey, chan)

