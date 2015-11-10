'''
Create Agenda
generate agenda from Asana tasks
and place them in Alpha Alpha Gdrive
folder to share with the chapter.

author(s):
	Jorge Rojas
'''
import sys
import math
import time
import json, base64
import urllib, urllib2
from docx import Document
from datetime import datetime
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

today = lambda: datetime.fromtimestamp( time.time() )
#Change '***************:' auth token from Asana once the members are registered
auth_token = base64.encodestring( '3IOlvKIK.qCLaoBd9o9vQzENWfspMQUl:' ).replace( '\n', '' )

heading = 'La Unidad Latina, Lambda Upsilon Lambda Fraternity, Inc.'
info = 'Alpha Alpha Chapter Agenda {0.year}-{0.month}-{0.day}'.format( today() )
attendance = '{0} - {1}'	#1 - Name; 2 - Status
topic_section = '{0} - mediator: {1}'	#discussion topic heading: 0 - topic number; 1 - topic name;

def curl( url, data = None, authToken = None ):
	if data is not None:
		req = urllib2.Request( url, data )
	else:
		req = urllib2.Request( url )

	if authToken is not None:
		req.add_header( 'Authorization', 'Basic %s'%authToken )

	response = urllib2.urlopen( req )
	res = response.read()
	return res

def getUser( taskStructDict ):
	user_id = taskStructDict[ 'assignee' ][ 'id' ]
	user_name = taskStructDict[ 'assignee' ][ 'name' ]

	user_url = 'https://app.asana.com/api/1.0/users/%s?opt_pretty'%user_id
	user = curl( user_url, authToken = auth_token )
	j_user = json.loads( user )

	return j_user[ 'data' ][ 'name' ], j_user[ 'data' ][ 'email' ]

def postLinkToSlack( gFileLink ):
	slack_webhook_url = 'https://hooks.slack.com/services/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'
	slack_msg = 'Hey <!everyone>, <%s|click this link> to see this week\'s agenda!'%gFileLink
	slack_res = curl( slack_webhook_url, '{"channel": "#debug", "username": "ReminderBot", "text":"' + slack_msg + '", "icon_emoji": ":mega:"}')

def uploadToDrive( fileName, projKey, gFolders ):
	'''Google Drive Authentication'''
	gauth = GoogleAuth()
	gauth.LocalWebserverAuth()
	drive = GoogleDrive(gauth)	#gdrive service instance

	parent_folder_id = ''
	for indx, cur_folder_name in enumerate(gFolders):
		if indx == 0:
			file1 = drive.ListFile({'q': "'root' in parents and trashed=false"}).GetList()
		else:
			file1 = drive.ListFile({'q': "'%s' in parents and trashed=false"%parent_folder_id}).GetList()
		for cur_file in file1:
			if cur_file['title'] == cur_folder_name:
				parent_folder_id = cur_file['id']

	try:
		if projKey == 'debug':
			drive_agenda = drive.CreateFile({'title': fileName, 'parents': [{"id": parent_folder_id}]})
		else:
			drive_agenda = drive.CreateFile({'title': fileName[8:], 'parents': [{"id": parent_folder_id}]})
		drive_agenda.SetContentFile( fileName )
		drive_agenda.Upload()
		print 'Agenda was successfully saved to Google Drive!'
	except:
		print 'Agenda was successfully created, but failed to upload to the drive!'

	g_alternate_link = ''
	files = drive.ListFile({'q': "'%s' in parents and trashed=false"%parent_folder_id}).GetList()
	
	for cur_file in files:
		if projKey == 'debug':
			if cur_file['title'] == fileName:
				g_alternate_link = cur_file['alternateLink']
		else:
			if cur_file['title'] == fileName[8:]:
				g_alternate_link = cur_file['alternateLink']

	postLinkToSlack( g_alternate_link )

def writeAgenda( authToken, projectUrl, document, fileName ):
	'''Capture all tasks (discussion topics)'''
	all_tasks = curl( projectUrl, authToken = authToken )
	j_all_tasks = json.loads( all_tasks )

	topic_cnt = 1
	isAttendance = False

	for cur_task in j_all_tasks[ 'data' ]:
		
		cur_task_name = cur_task[ 'name' ]
		cur_task_id = cur_task[ 'id' ]

		if cur_task_name.endswith( ':' ):
			if cur_task_name.find( 'Attendance') >= 0:
				isAttendance = True
			else:
				isAttendance = False
			document.add_heading( cur_task_name[:-1], level = 1 )
			topic_cnt = 1
		else:
			task_url = 'https://app.asana.com/api/1.0/tasks/%s'%cur_task_id
			story_url = 'https://app.asana.com/api/1.0/tasks/%s/stories'%cur_task_id

			'''Task Struct'''
			cur_task_struct = curl( task_url, authToken = authToken )
			j_cur_task_struct = json.loads( cur_task_struct )
			j_cur_task = j_cur_task_struct[ 'data' ]

			'''Stories Struct'''
			cur_task_stories = curl( story_url, authToken = authToken )
			j_cur_task_stories = json.loads( cur_task_stories )
			stories = j_cur_task_stories[ 'data' ]	#list of dictionaries 

			if not j_cur_task[ 'completed' ] and isAttendance:
				document.add_paragraph( attendance.format( j_cur_task[ 'assignee' ][ 'name' ], cur_task_name ), style = 'List Bullet')
			elif not j_cur_task[ 'completed' ] and not isAttendance:
				if j_cur_task['assignee'] != None and cur_task_name != None:
					document.add_heading( topic_section.format( cur_task_name.encode('utf-8'), j_cur_task[ 'assignee' ][ 'name' ].encode('utf-8') ).decode('utf-8'), level = 2)
				elif j_cur_task['assignee'] == None and cur_task_name != None:
					document.add_heading( topic_section.format( cur_task_name.encode('utf-8'), 'EBOARD' ).decode('utf-8'), level = 2)

				for cur_story in stories:
		
					tab_cnt = 0
					comments = cur_story[ 'text' ].encode('utf-8').split('\n')
					for cur_comment in comments:
						for cur_char in cur_comment:
							if cur_char != '_':
								break
							else:
								tab_cnt += 1

						comment = cur_comment[tab_cnt:].decode('utf-8')
						# print comment
						tab_cnt += 1
						if tab_cnt <= 1 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet')
						elif tab_cnt <= 3 and tab_cnt > 1 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet %s'%tab_cnt)
						elif tab_cnt > 3 and comment != None:
							document.add_paragraph( comment, style = 'List Bullet 3')
						tab_cnt = 0

			topic_cnt += 1

	document.save( fileName )

def createAgenda(projKey):

	proj_key = projKey
	gdrive_path = '1. Agendas & Minutes/2015 - 2016'.split('/')
	projects = {'debug': '54999242167362', 'agenda': '47058473474991'}
	project_url = 'https://app.asana.com/api/1.0/projects/%s/tasks'%projects[proj_key]

	'''Document Formatting Variables'''
	fi_name = 'agendas/lulaa_weekly_agenda_{0.year}-{0.month}-{0.day}.docx'.format( today() )
	if proj_key == 'debug':
		fi_name = 'lulaa_agenda_debug.docx'

	agenda_doc = Document()
	agenda_doc.add_heading( heading, 0 )
	agenda_doc.add_paragraph( info ).bold = True
	writeAgenda( auth_token, project_url, agenda_doc, fi_name )
	uploadToDrive( fi_name, proj_key, gdrive_path )

if __name__ == '__main__':
	createAgenda()



